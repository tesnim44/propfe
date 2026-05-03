from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ACCENT = RGBColor(33, 78, 96)
ACCENT_LIGHT = RGBColor(228, 239, 243)
TEXT = RGBColor(32, 37, 43)
MUTED = RGBColor(96, 104, 112)


PROJECT_TITLE = "IBlog"
PROJECT_SUBTITLE = "Plateforme web communautaire de blogging"
REPORT_DATE = "3 mai 2026"
FRAMEWORK_LABEL = "PRO-PFE"
SCOPE_NOTE = (
    "Le présent travail couvre la validation fonctionnelle, l'intégration et les cas "
    "limites automatisés. Les tests de sécurité offensive et de performance sont hors "
    "du périmètre de ce document et seront traités par d'autres membres de l'équipe."
)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, color="214E60", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def set_font(run, name, size, bold=False, color=TEXT, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    for style_name, font_name, size, color in [
        ("Title", "Cambria", 24, ACCENT),
        ("Heading 1", "Cambria", 16, ACCENT),
        ("Heading 2", "Cambria", 13, ACCENT),
        ("Heading 3", "Cambria", 11, ACCENT),
    ]:
        style = doc.styles[style_name]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:ascii"), font_name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_footer(section, label):
    section.footer.is_linked_to_previous = False
    footer_paragraph = section.footer.paragraphs[0]
    for child in list(footer_paragraph._p):
        footer_paragraph._p.remove(child)
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_paragraph.add_run(label)
    set_font(run, "Calibri", 9, color=MUTED)


def add_cover(doc, report_title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(12)
    run = p.add_run(FRAMEWORK_LABEL)
    set_font(run, "Cambria", 12, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(10)
    run = p.add_run(report_title)
    set_font(run, "Cambria", 24, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(20)
    run = p.add_run(subtitle)
    set_font(run, "Calibri", 13, color=MUTED)

    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider.space_after = Pt(20)
    set_paragraph_border(divider)

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = [
        ("Projet", f"{PROJECT_TITLE} - {PROJECT_SUBTITLE}"),
        ("Date du document", REPORT_DATE),
        ("Périmètre", "Validation fonctionnelle, intégration et cas limites"),
        ("Statut", "Version livrable"),
        ("Note", "Sécurité et performance traitées séparément par l'équipe"),
    ]
    for i, (label, value) in enumerate(rows):
        left = table.cell(i, 0)
        right = table.cell(i, 1)
        left.text = label
        right.text = value
        set_cell_shading(left, "E4EFF3")
        for cell in (left, right):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.space_before = Pt(2)
                paragraph.space_after = Pt(2)
                for run in paragraph.runs:
                    set_font(run, "Calibri", 10, color=TEXT)
        for run in left.paragraphs[0].runs:
            set_font(run, "Calibri", 10, bold=True, color=ACCENT)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    note.space_before = Pt(12)
    run = note.add_run(
        "Ce document est rédigé à partir du dépôt IBlog et des résultats de la suite "
        "de tests automatisés observables dans l'espace de travail."
    )
    set_font(run, "Calibri", 11, color=TEXT)

    doc.add_section(WD_SECTION_START.NEW_PAGE)


def add_toc(doc, title, items):
    heading = doc.add_paragraph(title, style="Heading 1")
    heading.paragraph_format.space_after = Pt(10)
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_font(run, "Calibri", 11, color=TEXT)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(5)
    return paragraph


def add_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        set_font(run, "Calibri", 11, bold=True, color=ACCENT)
    run = p.add_run(text)
    set_font(run, "Calibri", 11, color=TEXT)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_font(run, "Calibri", 11, color=TEXT)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_font(run, "Calibri", 11, color=TEXT)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(hdr[idx], "D9EAF0")
        for run in hdr[idx].paragraphs[0].runs:
            set_font(run, "Calibri", 10, bold=True, color=ACCENT)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[idx].paragraphs:
                paragraph.space_after = Pt(2)
                for run in paragraph.runs:
                    set_font(run, "Calibri", 10, color=TEXT)
    doc.add_paragraph()
    return table


def build_pvt(path: Path):
    doc = Document()
    style_document(doc)
    add_cover(
        doc,
        "Plan de Validation des Tests (PVT)",
        "Cadre de validation pour le projet IBlog",
    )
    add_footer(doc.sections[0], "PVT - IBlog")
    add_footer(doc.sections[1], "PVT - IBlog")

    add_toc(
        doc,
        "Sommaire",
        [
            "1. Résumé exécutif",
            "2. Objectifs du test",
            "3. Portée des tests",
            "4. Planning de validation",
            "5. Outils et environnement",
            "6. Stratégie et types de tests",
            "7. Jeux d'essai",
            "8. Critères de validation",
        ],
    )

    add_heading(doc, "Résumé exécutif")
    add_paragraph(
        doc,
        "Le présent PVT définit la stratégie de validation du projet IBlog, une "
        "application web PHP/MySQL orientée blogging communautaire. Le système couvre "
        "l'authentification, la publication d'articles, la recherche, la gestion des "
        "communautés, l'administration des utilisateurs et les interactions sociales "
        "autour des articles.",
    )
    add_paragraph(doc, SCOPE_NOTE)
    add_paragraph(
        doc,
        "Le dépôt fournit une suite PHPUnit active structurée en trois niveaux: "
        "tests unitaires, tests d'intégration et tests avancés. Le plan présenté ici "
        "formalise ce périmètre, les données d'essai, les critères de conformité et "
        "les livrables attendus pour la campagne de validation.",
    )

    add_heading(doc, "1. Objectifs du test")
    add_bullets(
        doc,
        [
            "Vérifier la conformité des fonctionnalités principales avec les comportements attendus.",
            "Détecter rapidement les régressions sur les flux critiques du backend et des API.",
            "Confirmer la robustesse du système face aux cas limites et aux entrées invalides.",
            "Établir une base de validation répétable avant livraison académique ou démonstration.",
        ],
    )

    add_heading(doc, "2. Portée des tests")
    add_heading(doc, "2.1 Fonctionnalités incluses", level=2)
    add_bullets(
        doc,
        [
            "Inscription, connexion, récupération de mot de passe et gestion de session.",
            "Gestion des utilisateurs et opérations d'administration.",
            "Création, mise à jour, suppression logique et lecture des articles.",
            "Commentaires, likes, sauvegarde d'articles et consultation des listes sauvegardées.",
            "Recherche d'articles et de profils avec modes Articles, People et All.",
            "Communautés, adhésion, fils de discussion premium et messagerie associée.",
            "Pagination, validation d'entrées et cohérence des réponses JSON.",
        ],
    )
    add_heading(doc, "2.2 Fonctionnalités exclues", level=2)
    add_bullets(
        doc,
        [
            "Campagnes offensives de sécurité: injection SQL, XSS, fuzzing ou audit de durcissement.",
            "Tests de charge, mesures de latence sous concurrence et benchmark de performance.",
            "Validation exhaustive cross-browser de l'interface utilisateur.",
            "Vérification en production réelle des flux SMTP ou d'une infrastructure externe.",
        ],
    )

    add_heading(doc, "3. Planning de validation")
    add_table(
        doc,
        ["Phase", "Activités prévues", "Période"],
        [
            ("Préparation", "Mise à plat du dépôt, revue du périmètre et des suites actives", "1 mai 2026"),
            ("Conception", "Structuration des scénarios, jeux d'essai et critères d'acceptation", "1-2 mai 2026"),
            ("Exécution", "Lancement des suites Unit, Integration et Advanced", "2-3 mai 2026"),
            ("Consolidation", "Analyse des résultats et rédaction des rapports", "3 mai 2026"),
        ],
    )

    add_heading(doc, "4. Outils et environnement")
    add_table(
        doc,
        ["Élément", "Description"],
        [
            ("Langage / runtime", "PHP 8.2.12"),
            ("Cadre d'exécution", "XAMPP en environnement local"),
            ("Base de données", "MySQL via phpMyAdmin"),
            ("Outil principal de test", "PHPUnit 11.5.55"),
            ("Références du dépôt", "phpunit.xml, tests/Unit, tests/Integration, tests/Advanced"),
            ("Observation complémentaire", "Navigateur + DevTools pour vérifications ponctuelles si nécessaire"),
        ],
    )

    add_heading(doc, "5. Stratégie et types de tests")
    add_heading(doc, "5.1 Tests unitaires", level=2)
    add_paragraph(
        doc,
        "Les tests unitaires ciblent les contrôleurs et repositories isolés. Ils valident "
        "les helpers d'authentification, les opérations CRUD de base, les adaptations de "
        "schéma et certains comportements de compatibilité hérités du dossier backend.",
    )
    add_heading(doc, "5.2 Tests d'intégration", level=2)
    add_paragraph(
        doc,
        "Les tests d'intégration vérifient les flux complets entre API, session, base de "
        "données et sérialisation de réponse. Le périmètre couvre l'authentification, la "
        "gestion des articles, l'administration, la recherche, la pagination et les communautés.",
    )
    add_heading(doc, "5.3 Tests avancés et cas limites", level=2)
    add_paragraph(
        doc,
        "Les tests avancés se concentrent sur les entrées invalides, les méthodes HTTP "
        "non autorisées, les identifiants erronés, les charges utiles malformées, la "
        "normalisation des chemins hérités et la stabilité générale des API.",
    )
    add_heading(doc, "5.4 Exemples de scénarios représentatifs", level=2)
    add_table(
        doc,
        ["ID", "Scénario", "Résultat attendu"],
        [
            ("TV-01", "Inscription valide", "Création du compte et ouverture de session"),
            ("TV-02", "Mot de passe faible à l'inscription", "Rejet avec message de validation"),
            ("TV-03", "Publication d'un article complet", "Article publié et visible dans le flux"),
            ("TV-04", "Modification d'un article d'un autre utilisateur", "Rejet avec contrôle d'accès"),
            ("TV-05", "Recherche par mot-clé d'article", "Retour des résultats pertinents"),
            ("TV-06", "Accès aux threads par un non-membre", "Refus d'accès"),
            ("TV-07", "Pagination sur une dernière page partielle", "Retour du sous-ensemble restant"),
            ("TV-08", "Action API inconnue", "Réponse d'erreur explicite et stable"),
        ],
    )

    add_heading(doc, "6. Jeux d'essai")
    add_table(
        doc,
        ["Jeu d'essai", "Description"],
        [
            ("Utilisateur invité", "Aucune session ouverte pour valider les refus d'accès."),
            ("Utilisateur gratuit", "Compte standard utilisé pour les tests fonctionnels courants."),
            ("Utilisateur premium", "Compte autorisé aux fonctionnalités premium des communautés."),
            ("Administrateur", "Compte avec isAdmin = 1 pour le périmètre back-office."),
            ("Article publié", "Article créé par un autre utilisateur pour les tests de lecture et de sauvegarde."),
            ("Brouillon personnel", "Article non publié servant aux tests de reprise et d'édition."),
            ("Communauté existante", "Communauté avec adhésion et discussions activées."),
            ("Jeton de réinitialisation", "Cas valide et cas expiré pour le mot de passe oublié."),
        ],
    )

    add_heading(doc, "7. Critères de validation")
    add_bullets(
        doc,
        [
            "Chaque test automatisé du périmètre actif doit se terminer sans échec bloquant.",
            "Les codes HTTP, messages d'erreur et écritures en base doivent correspondre au scénario attendu.",
            "Aucune erreur fatale PHP, corruption de données ou incohérence de session ne doit apparaître.",
            "Le rapport final doit citer les chiffres réels d'exécution et les limites du périmètre.",
        ],
    )
    add_paragraph(
        doc,
        "Livrables attendus: ce PVT, un rapport final d'exécution, la commande de lancement "
        "de la suite PHPUnit et les chiffres consolidés par famille de tests.",
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def build_final_report(path: Path):
    doc = Document()
    style_document(doc)
    add_cover(
        doc,
        "Rapport Final de Validation des Tests",
        "Exécution réelle de la campagne de validation IBlog",
    )
    add_footer(doc.sections[0], "Rapport final - IBlog")
    add_footer(doc.sections[1], "Rapport final - IBlog")

    add_toc(
        doc,
        "Sommaire",
        [
            "1. Introduction",
            "2. Rappel des objectifs et du périmètre",
            "3. Méthodologie d'exécution",
            "4. Exécution des tests et résultats",
            "5. Synthèse et recommandations",
            "Annexe A. Commandes et chiffres clés",
        ],
    )

    add_heading(doc, "1. Introduction")
    add_paragraph(
        doc,
        "Ce rapport présente les résultats observés lors de l'exécution effective de la "
        "suite de tests automatisés du projet IBlog. Il s'appuie sur les fichiers du dépôt, "
        "sur la configuration active de PHPUnit et sur une exécution confirmée localement le "
        "3 mai 2026.",
    )
    add_paragraph(doc, SCOPE_NOTE)

    add_heading(doc, "2. Rappel des objectifs et du périmètre")
    add_bullets(
        doc,
        [
            "Valider les flux critiques d'authentification, d'administration et de gestion de session.",
            "Confirmer les opérations de publication, modification, suppression logique et interaction sur les articles.",
            "Vérifier la recherche, la pagination, les communautés et les règles d'accès premium.",
            "Contrôler la robustesse des API face aux méthodes interdites et aux données invalides.",
        ],
    )

    add_heading(doc, "3. Méthodologie d'exécution")
    add_heading(doc, "3.1 Outils utilisés", level=2)
    add_bullets(
        doc,
        [
            "PHPUnit 11.5.55 pour l'automatisation de la campagne.",
            "PHP 8.2.12 comme runtime d'exécution.",
            "XAMPP et MySQL/phpMyAdmin comme environnement local de référence.",
            "Analyse du dépôt et des fichiers de tests pour relier chaque suite à son périmètre réel.",
        ],
    )
    add_heading(doc, "3.2 Environnement de test", level=2)
    add_table(
        doc,
        ["Paramètre", "Valeur"],
        [
            ("Machine d'exécution", "Poste local de développement"),
            ("Date d'exécution confirmée", REPORT_DATE),
            ("Configuration active", "phpunit.xml à la racine du dépôt"),
            ("Suites actives", "Unit, Integration, Advanced"),
            ("Suites non actives", "Aucune suite Performance active dans la configuration courante"),
        ],
    )
    add_heading(doc, "3.3 Données de test", level=2)
    add_paragraph(
        doc,
        "Les jeux d'essai couvrent plusieurs profils utilisateurs, des articles publiés et "
        "brouillons, des communautés existantes, des identifiants invalides, des tokens de "
        "réinitialisation et diverses charges utiles JSON destinées à éprouver les API.",
    )
    add_heading(doc, "3.4 Méthode", level=2)
    add_numbered(
        doc,
        [
            "Exécution de la suite complète via la commande vendor\\bin\\phpunit.bat.",
            "Exécution complémentaire des suites Unit, Integration et Advanced séparément pour isoler les statistiques.",
            "Consolidation des résultats et traduction en lecture académique structurée.",
        ],
    )

    add_heading(doc, "4. Exécution des tests et résultats")
    add_heading(doc, "4.1 Synthèse globale", level=2)
    add_table(
        doc,
        ["Indicateur", "Valeur observée"],
        [
            ("Statut global", "Réussi"),
            ("Total des tests", "84"),
            ("Total des assertions", "471"),
            ("Échecs", "0"),
            ("Erreurs", "0"),
            ("Date de confirmation", REPORT_DATE),
        ],
    )
    add_paragraph(
        doc,
        "La campagne automatisée s'est terminée sans échec. Les résultats confirment la "
        "stabilité du périmètre backend validé par PHPUnit dans l'état actuel du dépôt.",
    )

    add_heading(doc, "4.2 Résultats par suite", level=2)
    add_table(
        doc,
        ["Suite", "Tests", "Assertions", "Lecture métier", "Statut"],
        [
            ("Unit", "13", "143", "Contrôleurs et repositories isolés", "Réussi"),
            ("Integration", "37", "180", "Flux complets API, session, base et sérialisation", "Réussi"),
            ("Advanced", "34", "148", "Cas limites, méthodes invalides et régressions", "Réussi"),
        ],
    )

    add_heading(doc, "4.3 Résultats fonctionnels majeurs", level=2)
    add_table(
        doc,
        ["Domaine", "Scénarios couverts", "Constat"],
        [
            ("Authentification", "Inscription, connexion, doublon email, mot de passe faible, reset valide/expiré", "Comportements conformes"),
            ("Administration", "Accès invité refusé, création et mise à jour d'utilisateurs, auto-suppression bloquée", "Conforme"),
            ("Articles", "Publication, brouillon, édition propriétaire, suppression logique, commentaires, likes, sauvegarde", "Conforme"),
            ("Recherche", "Recherche article, mode People, absence de résultats, caractères spéciaux", "Conforme"),
            ("Communautés", "Adhésion, non-duplication, accès aux threads, limitation premium", "Conforme"),
            ("Pagination", "Première page, dernière page partielle, taille variable, liste vide", "Conforme"),
        ],
    )

    add_heading(doc, "4.4 Cas limites et robustesse", level=2)
    add_bullets(
        doc,
        [
            "Les requêtes GET interdites sur certaines API sont correctement rejetées.",
            "Les identifiants d'articles inconnus ou invalides ne provoquent pas de corruption de données.",
            "Les charges utiles de couverture malformées sont refusées.",
            "Les actions inconnues retournent une erreur explicite plutôt qu'un comportement silencieux.",
            "La recherche continue à se comporter proprement avec des caractères spéciaux ou une base indisponible.",
        ],
    )

    add_heading(doc, "5. Synthèse et recommandations")
    add_heading(doc, "5.1 Points forts", level=2)
    add_bullets(
        doc,
        [
            "La suite automatisée couvre un périmètre métier large et cohérent pour un projet PHP natif.",
            "Les scénarios critiques d'authentification, d'articles, de recherche et de communautés sont protégés par des tests répétables.",
            "Aucun défaut bloquant n'a été observé dans le périmètre exécuté le 3 mai 2026.",
        ],
    )
    add_heading(doc, "5.2 Limites du présent rapport", level=2)
    add_bullets(
        doc,
        [
            "Ce rapport ne constitue pas un audit de sécurité offensif.",
            "Aucune campagne de charge ou de mesure de performance sous concurrence n'est documentée ici.",
            "Les aspects purement visuels de l'interface n'ont pas fait l'objet d'une preuve exhaustive dans cette campagne automatisée.",
        ],
    )
    add_heading(doc, "5.3 Recommandations", level=2)
    add_numbered(
        doc,
        [
            "Conserver la suite PHPUnit comme socle de non-régression et l'exécuter à chaque livraison.",
            "Ajouter ultérieurement une suite Performance active si l'équipe souhaite industrialiser ce volet dans le dépôt.",
            "Compléter ce rapport par le livrable sécurité et le livrable performance rédigés par les autres membres.",
        ],
    )

    add_heading(doc, "Annexe A. Commandes et chiffres clés")
    add_paragraph(
        doc,
        "Commande globale exécutée: vendor\\bin\\phpunit.bat",
        bold_prefix="Commande: ",
    )
    add_paragraph(
        doc,
        "Suites détaillées exécutées séparément: --testsuite Unit, --testsuite Integration, --testsuite Advanced.",
        bold_prefix="Détail: ",
    )
    add_table(
        doc,
        ["Commande", "Résultat"],
        [
            ("vendor\\bin\\phpunit.bat", "OK - 84 tests, 471 assertions"),
            ("vendor\\bin\\phpunit.bat --testsuite Unit", "OK - 13 tests, 143 assertions"),
            ("vendor\\bin\\phpunit.bat --testsuite Integration", "OK - 37 tests, 180 assertions"),
            ("vendor\\bin\\phpunit.bat --testsuite Advanced", "OK - 34 tests, 148 assertions"),
        ],
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main():
    docs_dir = Path(__file__).resolve().parents[1] / "docs"
    build_pvt(docs_dir / "PVT_IBlog_FR.docx")
    build_final_report(docs_dir / "Rapport_Final_IBlog_FR.docx")


if __name__ == "__main__":
    main()
