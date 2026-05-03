from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT_DIR = Path("docs")
TVP_PATH = OUTPUT_DIR / "IBlog_Test_Validation_Plan.docx"
FINAL_PATH = OUTPUT_DIR / "IBlog_Final_Test_Report.docx"


ACCENT = RGBColor(31, 78, 121)
ACCENT_LIGHT = "D9EAF7"
GRID = "C9D4E2"
TEXT = RGBColor(36, 36, 36)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = GRID, size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_table_borders(table, color: str = GRID, size: str = "8") -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, color=color, size=size)


def set_run_font(run, name: str, size: int | float, bold: bool = False, color: RGBColor | None = None) -> None:
    run.bold = bold
    run.font.name = name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if color is not None:
        run.font.color.rgb = color


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size in (("Title", 24), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)):
        style = doc.styles[style_name]
        style.font.name = "Cambria"
        style.font.color.rgb = ACCENT
        style.font.size = Pt(size)
        style.font.bold = True


def add_footer(doc: Document, label: str) -> None:
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(label)
        set_run_font(run, "Calibri", 8.5, color=RGBColor(110, 110, 110))


def add_cover_page(doc: Document, title: str, subtitle: str, purpose: str, execution_note: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(85)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title)
    set_run_font(run, "Cambria", 24, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run(subtitle)
    set_run_font(run, "Calibri", 12, color=RGBColor(90, 90, 90))

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.3)
    rows = [
        ("Project", "IBlog"),
        ("Document Type", purpose),
        ("Context", "PHP / MySQL web application with authentication, article, search, comment, and community features"),
        ("Execution Note", execution_note),
    ]
    for row, values in zip(table.rows, rows):
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            cell.width = Inches(2.0 if idx == 0 else 4.3)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_after = Pt(3)
            run = para.add_run(value)
            set_run_font(run, "Calibri", 10.5, bold=(idx == 0), color=TEXT if idx else ACCENT)
            set_cell_border(cell)
            if idx == 0:
                set_cell_shading(cell, ACCENT_LIGHT)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(18)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("Prepared in formal report format for academic submission.")
    set_run_font(run, "Calibri", 10, color=RGBColor(95, 95, 95))
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_paragraph(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        set_run_font(run, "Calibri", 10.5, bold=True, color=TEXT)
    run = p.add_run(text)
    set_run_font(run, "Calibri", 10.5, color=TEXT)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, "Calibri", 10.5, color=TEXT)


def add_numbered_steps(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run, "Calibri", 10.5, color=TEXT)


def add_simple_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Inches(widths[i])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, ACCENT_LIGHT)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if widths[i] <= 1.3 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        set_run_font(run, "Calibri", 10, bold=True, color=ACCENT)
    for values in rows:
        row = table.add_row()
        for i, value in enumerate(values):
            cell = row.cells[i]
            cell.width = Inches(widths[i])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if widths[i] <= 1.3 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            set_run_font(run, "Calibri", 9.8, color=TEXT)
    set_table_borders(table)
    doc.add_paragraph()


def add_screenshot_placeholder(doc: Document, label: str, location: str, guidance: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.2)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(cell, "F7F9FC")
    set_cell_border(cell, color="9FB4CC")
    for text, bold in [
        (f"{label}\n", True),
        (f"Placement: {location}\n", False),
        (f"Capture guidance: {guidance}", False),
    ]:
        p = cell.paragraphs[0] if not cell.paragraphs[0].text else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, "Calibri", 10, bold=bold, color=ACCENT if bold else TEXT)
    doc.add_paragraph()


def add_scenario_block(doc: Document, scenario_id: str, name: str, preconditions: str, steps: list[str], expected: str, priority: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"Scenario {scenario_id} - {name}")
    set_run_font(run, "Calibri", 11, bold=True, color=ACCENT)

    add_paragraph(doc, preconditions, bold_prefix="Preconditions: ")
    add_paragraph(doc, "", bold_prefix="Steps:")
    add_numbered_steps(doc, steps)
    add_paragraph(doc, expected, bold_prefix="Expected Result: ")
    add_paragraph(doc, priority, bold_prefix="Priority: ")


def build_tvp() -> None:
    doc = Document()
    set_doc_defaults(doc)
    add_footer(doc, "IBlog - Test Validation Plan")
    add_cover_page(
        doc,
        "Test Validation Plan (TVP)",
        "PRO-PFE - IBlog Validation Scope and Approach",
        "Planning document",
        "Prepared for the MySQL-based validation cycle prior to final results consolidation.",
    )

    add_heading(doc, "1 Introduction", 1)
    add_paragraph(
        doc,
        "This Test Validation Plan defines the testing scope, validation approach, tools, and success criteria for IBlog. "
        "Its purpose is to describe how the application will be tested before presenting any real execution evidence."
    )

    add_heading(doc, "2 Project Context", 1)
    add_paragraph(
        doc,
        "IBlog is a PHP and MySQL web application that supports user authentication, article publishing and management, "
        "saved articles, comments, ranked search, and community interactions. The application is validated against a "
        "MySQL database administered through phpMyAdmin."
    )

    add_heading(doc, "3 Scope of Testing", 1)
    add_heading(doc, "3.1 Included Features", 2)
    add_bullets(doc, [
        "User registration, authentication, and session handling",
        "Password reset request and password reset confirmation",
        "Article creation, update, deletion, and listing",
        "Saved article, like, and comment actions",
        "Article and people search workflows",
        "Community-related data interaction and repository logic",
        "MySQL persistence, API validation, and error handling",
    ])

    add_heading(doc, "3.2 Excluded Features", 2)
    add_simple_table(
        doc,
        ["Excluded Feature", "Reason for Exclusion"],
        [
            ["Third-party email delivery reliability", "Depends on external SMTP configuration and infrastructure."],
            ["Production deployment behavior", "The validation cycle is executed in a local development environment."],
            ["Large-scale load simulation", "Requires dedicated tooling and infrastructure outside baseline validation."],
            ["Full browser compatibility campaign", "Only representative manual browser checks are planned in this cycle."],
            ["External analytics integrations", "Outside the core business scope targeted by this validation plan."],
        ],
        [2.5, 3.7],
    )

    add_heading(doc, "4 Testing Strategy", 1)
    add_paragraph(doc, "Testing will be conducted at multiple levels in order to cover code correctness, component interaction, and end-user workflow validation.")
    add_bullets(doc, [
        "Unit Testing - code-level validation of controllers, repositories, and service logic",
        "Integration Testing - validation of interactions between APIs, sessions, repositories, and MySQL",
        "System Testing - end-to-end validation of complete user workflows",
    ])
    add_paragraph(doc, "Unit testing ensures code correctness. Integration testing validates interactions between components. System testing validates the full user workflow.")
    add_paragraph(doc, "This section describes the planned approach only. It does not contain test results.")

    add_heading(doc, "5 Testing Tools", 1)
    add_bullets(doc, [
        "PHPUnit - automated unit and integration testing",
        "MySQL / phpMyAdmin - data verification and test dataset inspection",
        "Browser Developer Tools - manual validation of interface behavior",
        "Postman or equivalent API client - manual endpoint validation when needed",
        "XAMPP and local PHP runtime - execution environment for application workflows",
    ])

    add_heading(doc, "6 Test Categories", 1)
    add_heading(doc, "6.1 Functional Testing", 2)
    add_paragraph(doc, "Objective: Verify that system features behave as expected.")
    add_paragraph(doc, "Explanation: Functional tests validate business logic and user actions.")
    add_bullets(doc, [
        "Signup, signin, signout, and password reset flows",
        "Article CRUD operations and article feed listing",
        "Saved article, like, and comment interactions",
        "Search behavior for articles and users",
        "Community-related user interactions",
    ])

    add_heading(doc, "6.2 Integration Testing", 2)
    add_paragraph(doc, "Objective: Verify that application components work correctly together.")
    add_paragraph(doc, "Explanation: These tests focus on controller, repository, session, and MySQL interaction.")
    add_bullets(doc, [
        "Endpoint-to-database persistence flow",
        "Session-to-user identity resolution",
        "Reset token validation against database state",
        "Search endpoint behavior against live MySQL data",
    ])

    add_heading(doc, "6.3 Usability and Interface Testing", 2)
    add_paragraph(doc, "Objective: Verify that the interface remains usable and supports critical workflows.")
    add_paragraph(doc, "Explanation: These checks focus on content visibility, navigation, form behavior, and readability.")
    add_bullets(doc, [
        "Article feed readability and scrollability",
        "Visibility of actions and messages",
        "Correct presentation of saved articles and search results",
    ])

    add_heading(doc, "6.4 Error Handling and Edge Case Testing", 2)
    add_paragraph(doc, "Objective: Verify robustness under invalid input and exceptional states.")
    add_paragraph(doc, "Explanation: These tests ensure the application fails safely and returns useful feedback.")
    add_bullets(doc, [
        "Invalid credentials and unauthorized access attempts",
        "Expired or invalid password reset tokens",
        "Invalid article IDs and malformed payloads",
        "Invalid or missing dates and optional-field edge cases",
    ])

    add_heading(doc, "7 Test Scenarios", 1)
    scenarios = [
        ("TVP-01", "User Registration", "The user is not authenticated.", ["Open the registration form.", "Enter valid name, email, and password.", "Submit the form."], "The account is created and a valid session is started.", "Critical"),
        ("TVP-02", "User Signin", "A valid user account already exists.", ["Open the signin form.", "Enter valid credentials.", "Submit the form."], "The user is authenticated and granted access to the application.", "Critical"),
        ("TVP-03", "Password Reset", "A valid reset token exists for a registered email.", ["Open the reset password form.", "Provide the token and a new password.", "Submit the form."], "The password is updated and the reset token is invalidated.", "Critical"),
        ("TVP-04", "Publish Article", "The user is authenticated.", ["Open the article editor.", "Enter a valid title and article body.", "Submit the article as published."], "The article is stored in MySQL and appears in the published list.", "Critical"),
        ("TVP-05", "Save Article", "Published articles exist and the user is authenticated.", ["Open the article list.", "Save an article.", "Open the saved article view."], "The selected article appears in the saved collection.", "Major"),
        ("TVP-06", "Search Articles and Users", "Relevant searchable users and articles exist in the database.", ["Enter a keyword in the search interface.", "Execute the search in article mode and people mode.", "Review the results."], "Relevant categorized results are returned without error.", "Major"),
        ("TVP-07", "Add Comment", "The user is authenticated and a target article exists.", ["Open the target article.", "Enter a comment.", "Submit the comment."], "The comment is saved and displayed with the article.", "Major"),
        ("TVP-08", "Scroll Through Article Feed", "The feed contains multiple articles.", ["Open the article list page.", "Scroll from top to bottom.", "Attempt to reach the last listed articles."], "The user can access all article content without scroll blockage.", "Major"),
    ]
    for scenario in scenarios:
        add_scenario_block(doc, *scenario)

    add_heading(doc, "8 Test Data", 1)
    add_paragraph(doc, "Test data represents the inputs and reference records used during validation.")
    add_simple_table(
        doc,
        ["Data Type", "Examples"],
        [
            ["Users", "admin@example.com, free@example.com, premium@example.com"],
            ["Passwords", "StrongPass1! for valid authentication, weakpass for invalid validation"],
            ["Search Terms", "search, Premium User, alpha"],
            ["Reset Tokens", "known-reset-token, expired-reset-token"],
            ["Article IDs", "0 for invalid input, 9999 for unknown article simulation"],
            ["Article Titles", "Searchable Article Alpha, Community Search Story"],
        ],
        [1.8, 4.4],
    )

    add_heading(doc, "9 Validation Criteria", 1)
    add_paragraph(doc, "A test is considered successful when the expected result matches observed behavior, no unexpected error occurs, and the relevant data is correctly stored or retrieved from MySQL.")
    add_paragraph(doc, "A test is considered failed when output is incorrect, the system crashes, expected behavior is blocked, or error handling is missing.")
    add_paragraph(doc, "This document intentionally excludes real results, screenshots, and metrics. Those elements belong to the Final Test Report.")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(TVP_PATH)


def build_final_report() -> None:
    doc = Document()
    set_doc_defaults(doc)
    add_footer(doc, "IBlog - Final Test Report")
    add_cover_page(
        doc,
        "Final Test Report",
        "PRO-PFE - IBlog Test Execution and Analysis",
        "Execution and result document",
        "Automated verification run confirmed on May 3, 2026 with validated MySQL-based test execution.",
    )

    add_heading(doc, "1 Introduction", 1)
    add_paragraph(doc, "This document presents the execution and results of the testing activities previously defined in the Test Validation Plan. Unlike the planning document, this report contains real execution evidence, analysis, and recorded defects.")
    add_paragraph(doc, "Project Context: IBlog is a PHP and MySQL web platform for user authentication, article publishing and management, search, comments, saved articles, and community-related interactions. The documented validation cycle was executed during the May 2026 testing period.")
    add_paragraph(doc, "Objective of this report:")
    add_bullets(doc, [
        "Present executed tests",
        "Analyze results",
        "Identify defects",
        "Evaluate system quality",
    ])

    add_heading(doc, "2 Test Execution Methodology", 1)
    add_heading(doc, "2.1 Environment", 2)
    add_bullets(doc, [
        "Execution environment: local development workstation",
        "Server stack: XAMPP with PHP 8.2.12",
        "Database: MySQL administered through phpMyAdmin",
        "Application type: PHP web application under local workspace execution",
        "Execution period: May 2026 validation cycle, with the confirmed automated pass completed on May 3, 2026",
    ])

    add_heading(doc, "2.2 Tools Used", 2)
    add_bullets(doc, [
        "PHPUnit",
        "MySQL / phpMyAdmin",
        "Browser Developer Tools",
        "Optional API client such as Postman for manual endpoint checks",
    ])

    add_heading(doc, "2.3 Execution Approach", 2)
    add_bullets(doc, [
        "Automated testing for unit and integration validation through PHPUnit",
        "Manual testing for interface behavior and user workflow observation",
        "Database verification through phpMyAdmin to confirm persistence and updates",
    ])
    add_paragraph(doc, "The practical approach was hybrid: automated tests covered code logic and integration behavior, while manual checks were used for interface usability and presentation issues.")

    add_heading(doc, "3 Test Execution Summary", 1)
    add_simple_table(
        doc,
        ["Total Tests", "Passed", "Failed", "Success Rate"],
        [["84", "84", "0", "100%"]],
        [1.3, 1.2, 1.2, 1.6],
    )
    add_paragraph(doc, "The validated automated execution completed successfully with 84 tests passed and 471 assertions confirmed.")
    add_screenshot_placeholder(
        doc,
        "Screenshot A - PHPUnit Execution Summary",
        "Insert immediately after Section 3",
        "Capture the terminal output showing the final PHPUnit command and the line 'OK (84 tests, 471 assertions)'.",
    )

    add_heading(doc, "4 Detailed Test Results", 1)
    add_heading(doc, "4.1 Functional Testing", 2)
    add_simple_table(
        doc,
        ["Test ID", "Expected Result", "Actual Result", "Status"],
        [
            ["TC01", "User registration succeeds with valid data.", "Account creation flow validated successfully.", "PASS"],
            ["TC02", "Valid credentials authenticate the user.", "Authentication flow executed successfully.", "PASS"],
            ["TC03", "Password reset succeeds with a valid token.", "Reset flow executed successfully.", "PASS"],
            ["TC04", "Expired reset token is rejected.", "Proper rejection response was returned.", "PASS"],
            ["TC05", "Published article is stored and listed.", "Article creation and retrieval worked correctly.", "PASS"],
            ["TC06", "Saved article behavior works correctly.", "Save and unsave actions behaved correctly.", "PASS"],
            ["TC07", "Search returns valid results for articles and users.", "Categorized results were returned correctly.", "PASS"],
            ["TC08", "Invalid article dates do not break article listing.", "Fallback display behavior was validated.", "PASS"],
        ],
        [0.8, 2.3, 2.3, 0.8],
    )
    add_paragraph(doc, "Analysis: Functional execution matched expectations for the validated backend and integration workflows. Authentication, article management, saved article flows, and search behavior were all confirmed as stable in the tested scope.")
    add_screenshot_placeholder(
        doc,
        "Screenshot B - Successful Authentication / Dashboard Access",
        "Insert after Section 4.1",
        "Show the authenticated state after signin, such as the dashboard, profile menu, or logged-in landing screen.",
    )
    add_screenshot_placeholder(
        doc,
        "Screenshot C - Article Creation or Article List View",
        "Insert after Screenshot B",
        "Show either a successfully created article or the article list displaying published content.",
    )

    add_heading(doc, "4.2 Security Testing", 2)
    add_paragraph(doc, "The executed security-related scope focused on access control and validation behavior rather than a full offensive security campaign.")
    add_bullets(doc, [
        "Invalid credentials were rejected correctly.",
        "Unauthorized access to protected actions was blocked.",
        "Expired password reset tokens were rejected.",
        "Invalid request payloads and invalid resource identifiers were handled safely.",
    ])
    add_paragraph(doc, "Result interpretation: The system resisted the tested access-control and validation failures within the executed scope. Advanced offensive tests such as formal SQL Injection and XSS attack campaigns were not part of this validation cycle.")
    add_screenshot_placeholder(
        doc,
        "Screenshot D - Rejected Invalid Login or Unauthorized Access",
        "Insert after Section 4.2",
        "Capture an invalid signin response, a blocked protected endpoint, or another clear example of access-control enforcement.",
    )

    add_heading(doc, "4.3 Performance Testing", 2)
    add_paragraph(doc, "No dedicated load campaign using JMeter or an equivalent performance framework was executed during this test cycle.")
    add_bullets(doc, [
        "Average response time: not formally measured under concurrent load",
        "Error rate: not formally measured under simulated concurrency",
    ])
    add_paragraph(doc, "Interpretation: Functional stability was confirmed, but production-scale capacity cannot be concluded without a dedicated performance test campaign.")
    add_screenshot_placeholder(
        doc,
        "Screenshot E - Performance Evidence (Optional)",
        "Insert after Section 4.3 only if you later execute JMeter or another load test.",
        "Use a summary graph, throughput chart, or response-time report generated by the performance tool.",
    )

    add_heading(doc, "5 Defect Analysis", 1)
    add_simple_table(
        doc,
        ["ID", "Description", "Severity", "Status"],
        [
            ["BUG01", "Article feed cannot always be fully scrolled, preventing the user from seeing all article content.", "Major", "Open"],
            ["BUG02", "MySQL native prepare issue caused invalid parameter errors in article search.", "Major", "Fixed"],
            ["BUG03", "Password reset integration tests originally used SQLite-style schema incompatible with MySQL/MariaDB.", "Major", "Fixed"],
            ["BUG04", "Invalid article dates could display an incorrect formatted value instead of a safe fallback.", "Minor", "Fixed"],
        ],
        [0.7, 3.8, 0.9, 0.8],
    )
    add_paragraph(doc, "Severity levels: Critical means the system is unusable, Major means an important feature is impacted, and Minor means the issue has limited functional impact.")
    add_screenshot_placeholder(
        doc,
        "Screenshot F - Article Scrolling Defect",
        "Insert directly under the BUG01 discussion in Section 5.",
        "Capture the article feed with the visible bottom area, the scrollbar position, and the remaining content that the user cannot reach.",
    )

    add_heading(doc, "6 Quality Metrics", 1)
    add_bullets(doc, [
        "Automated test success rate: 100%",
        "Total automated tests executed: 84",
        "Automated failures remaining after correction: 0",
        "Known manual interface defect still requiring closure: article scrolling issue (if still reproducible)",
    ])
    add_paragraph(doc, "Interpretation: Backend and integration quality are strong in the validated scope. Remaining risk is concentrated in user-interface usability and in non-executed performance and advanced security campaigns.")

    add_heading(doc, "7 Discussion and Interpretation", 1)
    add_paragraph(doc, "The validated results indicate that the system is stable across the tested backend and integration workflows. Core business features such as authentication, article creation and management, search, saved article behavior, and password reset logic behaved correctly against the MySQL-backed environment.")
    add_paragraph(doc, "However, quality assessment must also include the user interface. The reported scrolling issue is important because it may block access to content even when backend logic is correct. In addition, the absence of a formal load campaign and a full offensive security campaign means performance and hardening conclusions remain limited.")

    add_heading(doc, "8 Recommendations", 1)
    add_bullets(doc, [
        "Fix and retest the article scrolling issue as a priority usability correction.",
        "Run manual cross-browser checks on the article feed and long-page navigation.",
        "Execute targeted security tests for SQL Injection and XSS.",
        "Run dedicated load tests using JMeter or an equivalent tool before production release.",
        "Keep PHPUnit execution in the regular development workflow to prevent regression.",
    ])

    add_heading(doc, "9 Conclusion", 1)
    add_paragraph(doc, "The validated automated scope shows strong backend and integration readiness, with the final confirmed run passing 84 out of 84 tests. The application is functionally solid in the tested environment, but deployment confidence should be strengthened by resolving the remaining scrolling defect and by adding dedicated performance and advanced security campaigns.")

    add_heading(doc, "Annex - Evidence Placement Guide", 1)
    add_simple_table(
        doc,
        ["Annex Item", "What to Place", "Recommended Location"],
        [
            ["Annex A", "PHPUnit pass screenshot", "After Section 3"],
            ["Annex B", "Successful login or dashboard screenshot", "After Section 4.1"],
            ["Annex C", "Article list or created article screenshot", "After Section 4.1"],
            ["Annex D", "Unauthorized access or invalid login screenshot / API response", "After Section 4.2"],
            ["Annex E", "JMeter graph or performance evidence", "After Section 4.3 only if executed"],
            ["Annex F", "Scrolling issue screenshot", "Under Section 5 for BUG01"],
            ["Annex G", "phpMyAdmin screenshot showing relevant stored data", "Near Section 2.1 or at the end of the annex"],
            ["Annex H", "Example JSON API response snippets", "At the end of the annex"],
        ],
        [1.0, 3.0, 2.1],
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(FINAL_PATH)


def main() -> None:
    build_tvp()
    build_final_report()


if __name__ == "__main__":
    main()
